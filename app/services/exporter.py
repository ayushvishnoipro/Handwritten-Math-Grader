"""
Document export service for generating Word documents from grading results.

Provides functionality to export grading results and analysis
to editable Word document format.
"""

from docx import Document
from docx.shared import Inches
from typing import Any, Dict
import json
from pathlib import Path
import io


def save_docx(content: str, output_path: str) -> None:
    """
    Save content to a Word document.
    
    Args:
        content: Text content to save
        output_path: Path where the document should be saved
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Handwritten Math Grader Report', 0)
    
    # Add content as paragraphs
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle different content types
        if line.startswith('==='):
            # Section headers
            doc.add_heading(line.replace('=', '').strip(), level=1)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.startswith('{') or line.startswith('['):
            # JSON content - format as code
            try:
                # Try to pretty-print JSON
                parsed = json.loads(line)
                formatted_json = json.dumps(parsed, indent=2)
                p = doc.add_paragraph()
                p.style = 'Intense Quote'
                p.add_run(formatted_json)
            except:
                # If not valid JSON, add as regular text
                doc.add_paragraph(line)
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_path)


def create_grading_report(grading_results: Dict[str, Any], 
                         output_path: str = "grading_report.docx") -> str:
    """
    Create a comprehensive grading report as a Word document.
    
    Args:
        grading_results: Complete grading results dictionary
        output_path: Path for the output document
        
    Returns:
        Path to the created document
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Math Solution Grading Report', 0)
    
    # Timestamp
    timestamp = grading_results.get('timestamp', 'Unknown')
    doc.add_paragraph(f'Generated: {timestamp}')
    doc.add_paragraph()  # Empty line
    
    # Overall Score Section
    doc.add_heading('Overall Results', level=1)
    overall_score = grading_results.get('overall_score', 0)
    p = doc.add_paragraph()
    p.add_run('Overall Score: ').bold = True
    p.add_run(f'{overall_score}/10')
    
    # Score interpretation
    if overall_score >= 8:
        interpretation = "Excellent work with strong mathematical reasoning"
    elif overall_score >= 6:
        interpretation = "Good work with minor areas for improvement"
    elif overall_score >= 4:
        interpretation = "Adequate work but needs significant improvement"
    else:
        interpretation = "Needs major revision and additional support"
    
    doc.add_paragraph(f'Interpretation: {interpretation}')
    doc.add_paragraph()
    
    # Symbolic Checks Section
    doc.add_heading('Mathematical Validation', level=1)
    symbolic_checks = grading_results.get('symbolic_checks', {})
    
    symbolic_score = symbolic_checks.get('score', 0)
    p = doc.add_paragraph()
    p.add_run('Symbolic Score: ').bold = True
    p.add_run(f'{symbolic_score}/10')
    
    # List individual checks
    checks = symbolic_checks.get('checks', [])
    if checks:
        doc.add_paragraph('Detailed Checks:')
        for check in checks:
            check_type = check.get('type', 'Unknown')
            is_correct = check.get('correct', False)
            message = check.get('message', 'No message')
            
            p = doc.add_paragraph()
            p.add_run('• ')
            p.add_run(f"{check_type}: ").bold = True
            
            if is_correct:
                p.add_run('✓ ').bold = True
                p.add_run(message)
            else:
                p.add_run('✗ ').bold = True
                p.add_run(message)
    
    # Geometry Results
    geometry_result = symbolic_checks.get('geometry_result')
    if geometry_result:
        doc.add_paragraph()
        doc.add_heading('Geometry Analysis', level=2)
        
        if geometry_result.get('success'):
            doc.add_paragraph('✓ Geometry problem solved successfully')
            interpretation = geometry_result.get('interpretation', '')
            if interpretation:
                doc.add_paragraph(f'Solution: {interpretation}')
        else:
            doc.add_paragraph('✗ Geometry problem could not be solved')
            error = geometry_result.get('error', '')
            if error:
                doc.add_paragraph(f'Error: {error}')
    
    doc.add_paragraph()
    
    # LLM Evaluation Section
    doc.add_heading('AI Evaluation', level=1)
    llm_evaluation = grading_results.get('llm_evaluation', {})
    
    llm_score = llm_evaluation.get('score', 0)
    p = doc.add_paragraph()
    p.add_run('AI Score: ').bold = True
    p.add_run(f'{llm_score}/10')
    
    # Feedback
    feedback = llm_evaluation.get('feedback', 'No feedback available')
    doc.add_paragraph()
    doc.add_paragraph('Feedback:')
    doc.add_paragraph(feedback)
    
    # Criteria breakdown
    criteria = llm_evaluation.get('criteria', {})
    if criteria:
        doc.add_paragraph()
        doc.add_heading('Evaluation Criteria', level=2)
        
        for criterion, result in criteria.items():
            score = result.get('score', 0)
            comment = result.get('comment', 'No comment')
            
            p = doc.add_paragraph()
            p.add_run(f'{criterion.title()}: ').bold = True
            p.add_run(f'{score}/10 - {comment}')
    
    # Suggestions
    suggestions = llm_evaluation.get('suggestions', [])
    if suggestions:
        doc.add_paragraph()
        doc.add_heading('Improvement Suggestions', level=2)
        for suggestion in suggestions:
            doc.add_paragraph(f'• {suggestion}')
    
    # OCR Results Section (Summary)
    doc.add_paragraph()
    doc.add_heading('OCR Processing Summary', level=1)
    
    ocr_results = grading_results.get('ocr_results', {})
    question_regions = ocr_results.get('question', [])
    solution_regions = ocr_results.get('solution', [])
    
    doc.add_paragraph(f'Question regions detected: {len(question_regions)}')
    doc.add_paragraph(f'Solution regions detected: {len(solution_regions)}')
    
    # Processed text
    processed_text = grading_results.get('processed_text', '')
    if processed_text:
        doc.add_paragraph()
        doc.add_heading('Extracted Solution Text', level=2)
        p = doc.add_paragraph(processed_text)
        p.style = 'Intense Quote'
    
    # Technical Details (Optional)
    if llm_evaluation.get('mock_evaluation'):
        doc.add_paragraph()
        doc.add_heading('Technical Notes', level=1)
        mock_reason = llm_evaluation.get('mock_reason', '')
        doc.add_paragraph(f'Note: This evaluation used a fallback system due to: {mock_reason}')
    
    # Save the document
    doc.save(output_path)
    return output_path


def export_to_json(grading_results: Dict[str, Any], output_path: str = "grading_results.json") -> str:
    """
    Export grading results to JSON format.
    
    Args:
        grading_results: Complete grading results dictionary
        output_path: Path for the output JSON file
        
    Returns:
        Path to the created JSON file
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(grading_results, f, indent=2, ensure_ascii=False)
    
    return output_path


def create_summary_report(grading_results: Dict[str, Any]) -> str:
    """
    Create a brief summary report as plain text.
    
    Args:
        grading_results: Complete grading results dictionary
        
    Returns:
        Summary text
    """
    overall_score = grading_results.get('overall_score', 0)
    symbolic_score = grading_results.get('symbolic_checks', {}).get('score', 0)
    llm_score = grading_results.get('llm_evaluation', {}).get('score', 0)
    
    summary = f"""
GRADING SUMMARY
===============

Overall Score: {overall_score}/10
- Mathematical Validation: {symbolic_score}/10
- AI Evaluation: {llm_score}/10

"""
    
    # Add key feedback
    feedback = grading_results.get('llm_evaluation', {}).get('feedback', '')
    if feedback:
        summary += f"Key Feedback: {feedback[:200]}...\n\n"
    
    # Add main issues/successes
    checks = grading_results.get('symbolic_checks', {}).get('checks', [])
    correct_checks = [c for c in checks if c.get('correct')]
    incorrect_checks = [c for c in checks if not c.get('correct')]
    
    if correct_checks:
        summary += f"Strengths: {len(correct_checks)} validation checks passed\n"
    
    if incorrect_checks:
        summary += f"Areas for improvement: {len(incorrect_checks)} issues identified\n"
    
    return summary
